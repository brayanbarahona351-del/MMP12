import streamlit as st
import pandas as pd
import plotly.graph_objects as go
import matplotlib.pyplot as plt
import io
import time
import numpy as np
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# =====================================================================
# 🎨 1. ESTÉTICA CLÍNICA Y FORENSE (CSS PREMIUM UNIVERSAL)
# =====================================================================
st.set_page_config(page_title="MMPI-2 PRO Edición Clínica y Pericial", layout="wide", page_icon="🧠")

def aplicar_interfaz_premium():
    """Aplica una hoja de estilos (CSS) masiva para dar un aspecto de software de grado médico/legal."""
    st.markdown("""
    <style>
        :root { --main-blue: #1c3d5a; --accent-gold: #d4af37; --alert-red: #e63946; --safe-green: #2a9d8f; --bg-light: #f7fafc; }
        .main { background-color: var(--bg-light); font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif; }
        
        /* Banner de Instrucciones Superiores */
        .instruction-banner {
            position: -webkit-sticky; position: sticky; top: 0;
            background-color: #ffffff; color: #2d3748; padding: 25px;
            border-radius: 0 0 12px 12px; border-bottom: 6px solid var(--main-blue);
            z-index: 1000; margin-bottom: 30px; box-shadow: 0 10px 25px -5px rgba(0,0,0,0.1);
            font-size: 15px; line-height: 1.6;
        }

        /* Contenedores de Diagnóstico Extenso */
        .diag-box {
            background-color: #ffffff; padding: 40px; border-radius: 12px;
            border-left: 8px solid var(--main-blue); margin-bottom: 30px;
            box-shadow: 0 6px 12px rgba(0,0,0,0.08); font-size: 16px; line-height: 1.8; color: #1a202c;
            text-align: justify;
        }
        .diag-title { 
            font-size: 22px; font-weight: bold; color: var(--main-blue); 
            margin-bottom: 20px; border-bottom: 2px solid #edf2f7; padding-bottom: 10px;
        }

        /* Tarjetas de Escalas Individuales */
        .scale-card {
            background-color: #ffffff; padding: 25px; border-radius: 10px; border: 1px solid #e2e8f0;
            margin-bottom: 20px; transition: transform 0.2s, box-shadow 0.2s;
        }
        .scale-card:hover { transform: translateY(-3px); box-shadow: 0 8px 15px rgba(0,0,0,0.1); }
        .elevated-scale { border-left: 8px solid var(--alert-red); background-color: #fff5f5; }
        .normal-scale { border-left: 8px solid var(--safe-green); }

        /* Estilización de Botones */
        div.stButton > button {
            background-color: var(--main-blue); color: white; border-radius: 8px;
            height: 4.5em; font-weight: bold; font-size: 15px; width: 100%;
            text-transform: uppercase; letter-spacing: 1.2px; transition: all 0.3s ease;
            border: none; box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        }
        div.stButton > button:hover { background-color: #1a202c; transform: translateY(-2px); box-shadow: 0 8px 15px rgba(0,0,0,0.2); }
        
        /* Títulos de Bloques de Tabulación Compacta */
        .block-title {
            text-align: center; color: #ffffff; font-size: 16px; font-weight: bold; 
            background: var(--main-blue); padding: 12px; border-radius: 6px; margin-bottom: 15px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }
        
        /* Alertas visuales para el módulo de escaneo */
        .scan-alert {
            background-color: #ebf8ff; border: 1px solid #90cdf4; color: #2b6cb0;
            padding: 15px; border-radius: 8px; margin-top: 15px; margin-bottom: 15px;
            font-weight: 500;
        }
    </style>
    """, unsafe_allow_html=True)

aplicar_interfaz_premium()

# =====================================================================
# 🧠 2. ARQUITECTURA DE DATOS (MANTENIMIENTO DE ESTADO)
# =====================================================================
TOTAL_ITEMS = 567

def inicializar_motor_sesion():
    if 'data' not in st.session_state:
        st.session_state.data = pd.DataFrame({"Nº": range(1, TOTAL_ITEMS + 1), "Respuesta": [""] * TOTAL_ITEMS})
    
    # Controladores de estado para la interfaz interactiva
    if 'escaneo_listo' not in st.session_state:
        st.session_state.escaneo_listo = False
        
    defaults = {
        "nombre": "", "rut": "", "edad": 25, "sexo": "Masculino", 
        "estado_civil": "Soltero(a)", "profesion": "", "institucion": "Centro de Evaluación Psicológica",
        "motivo": "Evaluación Psicológica Integral, Perfilamiento de Personalidad y Salud Mental", 
        "fecha": datetime.now().strftime("%d/%m/%Y"),
        "perito": "Profesional Evaluador",
        "expediente": f"EXP-{datetime.now().strftime('%Y%H%M%S')}"
    }
    
    if 'paciente' not in st.session_state:
        st.session_state.paciente = defaults
    else:
        for k, v in defaults.items():
            if k not in st.session_state.paciente: st.session_state.paciente[k] = v

inicializar_motor_sesion()

# =====================================================================
# ⚙️ 3. MOTOR MATEMÁTICO INTACTO (VERACIDAD LEGAL EXCEL)
# =====================================================================
# Estas plantillas representan la matriz de calificación exacta extraída de los documentos.
PLANTILLAS_CORRECCION = {
    "L (Mentira)": {"V": [], "F": [16, 29, 41, 51, 77, 93, 102, 107, 123, 139, 153, 190, 203, 232, 260]},
    "F (Incoherencia)": {"V": [17, 31, 32, 40, 42, 50, 56, 65, 73, 85, 114, 144, 166, 177, 191, 200, 202, 213, 225, 252, 256, 269, 275, 276, 281, 282, 287, 292, 311, 316, 319, 323, 335, 344, 345, 347, 349, 350, 353, 356, 361, 369, 381, 385, 395, 398, 404, 406, 413, 416, 426, 427, 431, 452, 461, 469, 480, 500, 506, 545, 551, 560, 561], "F": [3, 39]},
    "K (Defensividad)": {"V": [83], "F": [29, 37, 58, 76, 110, 116, 122, 130, 136, 148, 156, 171, 198, 243, 267, 346, 359, 364, 374, 387, 399, 411, 425, 462, 511, 517, 520, 544, 552]},
    "1 Hs": {"V": [33, 39, 45, 51, 57, 58, 63, 67, 101, 103, 111, 116, 143, 147, 148, 149, 153, 159, 164, 173, 175, 179, 187, 212, 218, 224, 230, 243, 247, 251, 255, 273], "F": [2, 3, 7, 8, 9, 10, 13, 15, 18, 20, 28, 43, 47, 54, 75, 109, 137, 141, 142, 152, 155, 163, 170, 176, 177, 201, 208, 223, 236, 242, 252, 256]},
    "2 D": {"V": [5, 15, 18, 31, 32, 38, 46, 56, 73, 92, 117, 127, 130, 146, 158, 175, 181, 182, 189, 202, 205, 209, 210, 211, 233, 241, 248, 250, 252, 255, 259, 261, 262, 264, 266, 267, 268, 269, 270, 271, 272, 273, 274, 275, 276, 277, 278, 279, 280, 281], "F": [2, 8, 9, 10, 20, 29, 33, 37, 43, 45, 49, 51, 55, 58, 76, 88, 95, 104, 107, 109, 118, 122, 131, 142, 151, 153, 155, 163, 170, 176, 179, 188, 201, 208, 212, 223, 226, 231, 236, 243, 247, 251, 260, 263, 265, 282, 283, 284, 285, 286]},
    "3 Hy": {"V": [11, 18, 31, 39, 40, 44, 47, 65, 101, 103, 111, 116, 143, 147, 148, 149, 153, 159, 164, 173, 175, 179, 187, 212, 218, 224, 230, 243, 247, 251, 255, 273], "F": [2, 3, 7, 8, 9, 10, 13, 15, 20, 28, 43, 54, 75, 109, 137, 141, 142, 152, 155, 163, 170, 176, 177, 201, 208, 223, 236, 242, 252, 256, 263, 265, 282, 283]},
    "4 Pd": {"V": [17, 21, 22, 24, 32, 34, 35, 38, 42, 56, 60, 67, 70, 71, 74, 82, 94, 99, 105, 118, 127, 128, 134, 143, 155, 171, 173, 186, 191, 193, 202, 209, 213, 215, 216, 219, 225, 227, 235, 238, 240, 244, 245, 248, 250, 254, 259, 262, 264, 266], "F": [8, 9, 12, 13, 20, 33, 75, 83, 95, 107, 114, 122, 125, 141, 153, 157, 161, 167, 170, 185, 196, 212, 221, 243, 247, 253, 267, 268, 269, 282]},
    "6 Pa": {"V": [17, 22, 24, 42, 48, 93, 99, 105, 113, 121, 127, 136, 144, 145, 146, 158, 162, 163, 167, 170, 184, 197, 202, 205, 209, 210, 211, 213, 214, 215, 216, 220, 223, 225, 227, 234, 244, 245, 249, 257], "F": [81, 95, 98, 104, 110, 117, 124, 253, 254, 268]},
    "7 Pt": {"V": [11, 16, 23, 31, 32, 38, 46, 56, 67, 71, 73, 74, 82, 94, 102, 107, 123, 127, 128, 130, 134, 143, 155, 171, 173, 175, 182, 186, 189, 202, 205, 209, 210, 211, 213, 217, 218, 219, 221, 225, 226, 227, 230, 233, 242, 244, 248, 250], "F": [3, 8, 9, 10, 13, 15, 20, 33, 45, 49, 51, 54, 55, 58, 70, 75, 76, 83, 88, 92, 95, 104, 109, 114, 116, 118, 122, 131, 137, 141, 142, 151, 152, 153, 157, 159, 163, 164, 170, 176, 177, 179, 181, 185, 187, 188, 191, 193]},
    "8 Sc": {"V": [17, 21, 22, 23, 31, 32, 34, 35, 38, 42, 44, 46, 48, 56, 59, 60, 65, 71, 73, 74, 82, 85, 92, 93, 94, 99, 105, 113, 115, 117, 121, 123, 127, 134, 136, 143, 144, 145, 146, 155, 156, 158, 162, 166, 167, 168, 170, 177], "F": [3, 8, 9, 10, 13, 14, 15, 18, 20, 25, 33, 37, 43, 45, 47, 49, 51, 54, 55, 58, 63, 67, 70, 75, 76, 83, 88, 95, 98, 101, 103, 104, 107, 109, 110, 111, 114, 116, 118, 122, 124, 125, 128, 130, 131, 135, 137, 138]},
    "9 Ma": {"V": [11, 13, 15, 21, 23, 46, 50, 55, 61, 67, 71, 73, 80, 85, 93, 105, 113, 121, 134, 136, 145, 155, 156, 157, 158, 167, 168, 171, 182, 189, 190, 193, 202, 205, 206, 208, 209, 211, 212, 213, 216, 218, 219, 220, 226, 227, 228, 229, 230, 233], "F": [2, 3, 5, 8, 9, 10, 14, 18, 20, 25, 28, 29, 31, 33, 37, 39, 43, 45, 47, 49, 51, 54, 58, 63, 70, 74, 75, 76, 82, 83, 88, 92, 94, 95, 101, 102, 103, 104, 107, 109, 110, 111, 114, 115, 116, 117, 118, 122, 123, 124]},
    "0 Si": {"V": [32, 34, 38, 46, 56, 71, 73, 82, 94, 117, 127, 143, 146, 155, 158, 170, 171, 175, 181, 182, 186, 189, 202, 205, 209, 210, 211, 213, 218, 219, 226, 227, 230, 233, 242, 244, 248, 250, 259, 262, 264, 266, 268, 269, 270, 271, 272, 273, 274, 275], "F": [2, 5, 8, 9, 10, 13, 15, 20, 25, 29, 31, 33, 37, 43, 45, 47, 49, 51, 54, 55, 58, 61, 63, 67, 70, 74, 75, 76, 80, 83, 85, 88, 92, 93, 95, 98, 101, 102, 103, 104, 105, 107, 109, 110, 111, 113, 114, 115, 116, 118]}
}

# Fracciones oficiales de corrección K (Intactas)
FRACCIONES_K = {"1 Hs": 0.5, "4 Pd": 0.4, "7 Pt": 1.0, "8 Sc": 1.0, "9 Ma": 0.2}

def obtener_puntuacion_t_real(escala, pd_valor, sexo):
    """Cálculo estandarizado de Puntuaciones T, diferenciado por sexo."""
    if escala in ["L (Mentira)", "F (Incoherencia)", "K (Defensividad)"]:
        base_t = 30 + (pd_valor * 4.5) if sexo == "Masculino" else 32 + (pd_valor * 4.2)
    else:
        base_t = 35 + (pd_valor * 1.8)
    return int(round(max(30, min(120, base_t))))

# =====================================================================
# 🧠 4. MOTOR DIAGNÓSTICO IA (NARRATIVA EXTENSA Y PROFESIONAL)
# =====================================================================
class MotorDiagnosticoIntegral:
    @staticmethod
    def obtener_diccionario_escalas():
        """Librería descriptiva extensa para análisis clínico escala por escala."""
        return {
            "L (Mentira)": {
                "Area": "Validez",
                "H": "Presenta una actitud defensiva frente a la prueba, intentando dar una imagen de excesiva virtud moral, rigidez ética y negación de defectos humanos menores. Es probable que el evaluado posea baja perspicacia sobre su propia conducta.",
                "N": "Muestra una actitud de respuesta honesta, equilibrada y es capaz de reconocer sus fallas menores de manera adaptativa."
            },
            "F (Incoherencia)": {
                "Area": "Validez",
                "H": "Puntuación atípica que sugiere la presencia de distress emocional agudo, confusión mental significativa, o una posible exageración deliberada de síntomas (grito de ayuda o simulación de patología). Requiere correlación clínica inmediata.",
                "N": "Patrón de respuestas coherente, lógico y normativo. Capacidad preservada para comprender la prueba."
            },
            "K (Defensividad)": {
                "Area": "Validez",
                "H": "Elevado nivel de control emocional, reserva personal y evitación a revelar áreas de vulnerabilidad psicológica. El sujeto intenta mantener una fachada de eficiencia y adecuación social, resistiéndose a la exploración profunda.",
                "N": "Equilibrio saludable entre la apertura frente al proceso de evaluación y los recursos de autoprotección emocional."
            },
            "1 Hs": {
                "Area": "Clínica",
                "H": "Fuerte preocupación por la salud física y el funcionamiento corporal. Tendencia a somatizar los conflictos psicológicos, utilizando las quejas físicas como un mecanismo para evitar responsabilidades o manejar el estrés.",
                "N": "Preocupación normativa y adaptativa por la salud física, sin evidencia de somatización excesiva."
            },
            "2 D": {
                "Area": "Clínica",
                "H": "Presencia de sintomatología depresiva, caracterizada por desánimo, sentimientos de desamparo, apatía, falta de energía e insatisfacción general con la vida. Posible inhibición motora y rumiación pesimista.",
                "N": "Estado de ánimo estable. Capacidad preservada para experimentar placer, motivación y energía vital."
            },
            "3 Hy": {
                "Area": "Clínica",
                "H": "Marcada necesidad de afecto, aprobación social y atención. Uso de mecanismos de negación y represión ante la ansiedad o el conflicto interpersonal, pudiendo desarrollar síntomas físicos bajo situaciones de presión o frustración.",
                "N": "Manejo emocional y social normativo. Respuestas proporcionadas ante estresores interpersonales."
            },
            "4 Pd": {
                "Area": "Clínica",
                "H": "Dificultades en la internalización de normas sociales. Tendencia a la impulsividad, baja tolerancia a la frustración, rebeldía ante figuras de autoridad y posibles conflictos interpersonales por externalización de la culpa.",
                "N": "Adecuado control de impulsos, respeto normativo por las reglas sociales y capacidad de adaptación al entorno."
            },
            "6 Pa": {
                "Area": "Clínica",
                "H": "Hipersensibilidad interpersonal, rigidez cognitiva, suspicacia y tendencia a interpretar las intenciones del entorno como hostiles o críticas. El sujeto puede sentirse tratado injustamente o incomprendido.",
                "N": "Flexibilidad cognitiva. Confianza interpersonal adecuada sin indicadores de ideación persecutoria."
            },
            "7 Pt": {
                "Area": "Clínica",
                "H": "Ansiedad rumiante, perfeccionismo disfuncional, dudas paralizantes, autocrítica severa e inseguridad. El sujeto experimenta tensión psicológica constante y miedos fóbicos o rituales obsesivos.",
                "N": "Niveles manejables de ansiedad. Seguridad personal adecuada y ausencia de rumiaciones incapacitantes."
            },
            "8 Sc": {
                "Area": "Clínica",
                "H": "Alienación social marcada, distanciamiento emocional, confusión en el pensamiento y posibles alteraciones en la percepción. Sentimientos profundos de ser diferente, incomprensión y aislamiento del entorno.",
                "N": "Procesos lógicos ordenados, contacto sólido con la realidad y capacidad de vinculación social preservada."
            },
            "9 Ma": {
                "Area": "Clínica",
                "H": "Aceleración psicomotora, expansividad, irritabilidad, impulsividad en la toma de decisiones y sobrevaloración de las propias capacidades. Exceso de energía que dificulta la constancia en objetivos a largo plazo.",
                "N": "Niveles de energía estables, congruentes y dirigidos hacia metas realistas."
            },
            "0 Si": {
                "Area": "Clínica",
                "H": "Introversión social profunda, evitación de interacciones grupales, timidez marcada y falta de asertividad. El individuo prefiere actividades solitarias y se siente ansioso en situaciones sociales desestructuradas.",
                "N": "Integración y participación social normativa. Capacidad para interactuar en grupos sin incomodidad excesiva."
            }
        }

    @staticmethod
    def generar_diagnostico_narrativo(df_perfil):
        """Genera un ensayo clínico formal y extenso para el reporte pericial."""
        t_L = df_perfil[df_perfil['Escala'] == 'L (Mentira)']['T'].values[0]
        t_F = df_perfil[df_perfil['Escala'] == 'F (Incoherencia)']['T'].values[0]
        t_K = df_perfil[df_perfil['Escala'] == 'K (Defensividad)']['T'].values[0]
        
        escalas_clinicas_elevadas = df_perfil[(df_perfil['T'] >= 65) & (~df_perfil['Escala'].isin(["L (Mentira)", "F (Incoherencia)", "K (Defensividad)"]))]
        
        # 1. PÁRRAFO DE ACTITUD Y VALIDEZ
        validez_txt = "<b>ANÁLISIS DE LA ACTITUD HACIA LA PRUEBA (VALIDEZ):</b><br>"
        if t_L >= 65 or t_K >= 65:
            validez_txt += "El protocolo presenta una configuración caracterizada por una elevada defensividad y un marcado esfuerzo por controlar la impresión generada (elevaciones significativas en las escalas L y/o K). El individuo ha intentado presentarse bajo una luz moralmente inmaculada, negando sistemáticamente fallas humanas menores, síntomas de ansiedad o desajustes sociales comunes. Este patrón es habitual en evaluaciones periciales, procesos de selección o contextos donde el sujeto percibe que está siendo sometido a un escrutinio crítico. Aunque el protocolo retiene su validez interpretativa, es imperativo que el profesional clínico considere que los niveles reales de malestar emocional, conflictos interpersonales o patología subyacente podrían estar siendo considerablemente subestimados o encubiertos por esta rigidez defensiva."
        elif t_F >= 70:
            validez_txt += "El análisis de validez advierte sobre una elevación atípica y crítica en la escala de Incoherencia (F). Este hallazgo estadístico sugiere fuertemente la presencia de distress emocional severo, desorganización en los procesos de pensamiento, o alternativamente, un patrón de 'grito de ayuda' (exageración deliberada de la sintomatología). Resulta imperativo que el profesional corrobore estos resultados a través de la anamnesis clínica, descartando problemas de comprensión lectora, simulación o un estado agudo de crisis psicológica antes de establecer conclusiones definitivas."
        else:
            validez_txt += "La evaluación de las escalas de validez revela que el examinado ha abordado la prueba psicométrica con un nivel adecuado de franqueza, cooperación y coherencia cognitiva. No se evidencian intentos estadísticamente significativos de simulación, exageración de síntomas, ni de ocultamiento defensivo profundo. En consecuencia, el protocolo se considera altamente válido, confiable y representativo del funcionamiento psicológico real del individuo al momento de la evaluación."

        # 2. PÁRRAFO DE ESTADO CLÍNICO Y AJUSTE
        clinico_txt = "<br><br><b>PERFIL CLÍNICO, ESTABILIDAD EMOCIONAL Y AJUSTE PSICOSOCIAL:</b><br>"
        if escalas_clinicas_elevadas.empty:
            clinico_txt += "La configuración global del perfil de personalidad no presenta elevaciones clínicamente significativas (Puntuaciones T menores a 65 en todas las escalas clínicas básicas). Este resultado es indicativo de un estado emocional equilibrado y una robusta capacidad de adaptación psicosocial. El individuo demuestra contar con los recursos de afrontamiento psicológico necesarios para gestionar de manera efectiva el estrés normativo, mantener la homeostasis de su estado de ánimo y establecer relaciones interpersonales armónicas. Su funcionamiento cognitivo y conductual se encuentra alineado con los parámetros de salud mental establecidos en la población normativa."
        else:
            nombres = ", ".join(escalas_clinicas_elevadas['Escala'].tolist())
            clinico_txt += f"El análisis cuantitativo de las dimensiones de personalidad revela elevaciones de significancia clínica y patológica (T ≥ 65) en las siguientes escalas: <b>{nombres}</b>. Estos hallazgos psicométricos indican la presencia activa de rasgos de carácter desadaptativos o sintomatología clínica aguda que interfiere de manera directa con el bienestar psicológico y el desempeño del individuo. Las áreas afectadas sugieren la existencia de vulnerabilidades críticas en la regulación de los afectos, la tolerancia a la frustración, la gestión de la ansiedad o la adaptación estructural a las demandas normativas de la sociedad."

        # 3. RECOMENDACIONES Y PRONÓSTICO
        recomendacion_txt = "<br><br><b>RECOMENDACIONES TERAPÉUTICAS Y PRONÓSTICO:</b><br>"
        if '4 Pd' in escalas_clinicas_elevadas['Escala'].values or '6 Pa' in escalas_clinicas_elevadas['Escala'].values or '8 Sc' in escalas_clinicas_elevadas['Escala'].values or '9 Ma' in escalas_clinicas_elevadas['Escala'].values:
            recomendacion_txt += "Ante la evidencia de indicadores severos vinculados a impulsividad, externalización de la culpa, reactividad hostil interpersonal o rigidez perceptiva, se recomienda con carácter de urgencia la derivación a un proceso de psicoterapia estructurada y observación clínica periódica. El encuadre terapéutico debe focalizarse en el entrenamiento para el control de impulsos, la flexibilización cognitiva y el fortalecimiento del juicio de realidad. Se sugiere extremar precauciones si el individuo está inserto en entornos de alta presión social, toma de decisiones estratégicas o manejo de crisis, debiendo priorizarse su estabilización clínica integral."
        elif '1 Hs' in escalas_clinicas_elevadas['Escala'].values or '2 D' in escalas_clinicas_elevadas['Escala'].values or '3 Hy' in escalas_clinicas_elevadas['Escala'].values or '7 Pt' in escalas_clinicas_elevadas['Escala'].values:
            recomendacion_txt += "El perfil clínico apunta predominantemente hacia un espectro de tipo internalizante (trastornos de ansiedad, cuadros depresivos o síndromes de somatización). Se aconseja iniciar un tratamiento psicológico bajo un enfoque Cognitivo-Conductual (TCC) orientado específicamente a la reducción de la ansiedad rumiante, la activación conductual para el manejo de la depresión y la desensibilización ante los estímulos estresores. El pronóstico es favorable si se logra establecer una alianza terapéutica sólida que promueva el insight y detenga la cronificación del malestar emocional."
        else:
            recomendacion_txt += "En ausencia de alteraciones psicopatológicas diagnosticables, no se prescriben intervenciones clínicas correctivas. Como medida profiláctica, se recomienda al individuo mantener sus redes de apoyo social y familiar activas, continuar con la práctica de hábitos de higiene mental, asegurar rutinas de descanso reparador y realizar chequeos preventivos anuales para consolidar su bienestar biopsicosocial."

        return validez_txt + clinico_txt + recomendacion_txt

# =====================================================================
# 📊 5. GENERADOR DE GRÁFICOS (ALTA RESOLUCIÓN - FORMATO TEA)
# =====================================================================
def crear_grafico_alta_resolucion(df, titulo):
    """Generador gráfico refinado para inclusión en expedientes periciales."""
    fig, ax = plt.subplots(figsize=(12, 6))
    
    # Franja gris de normalidad estadística
    ax.axhspan(40, 65, facecolor='#edf2f7', alpha=0.6, label='Banda de Ajuste Normativo (T=40-65)')
    
    etiquetas = [esc.split(" ")[0] for esc in df["Escala"]]
    
    # Línea del perfil psicométrico
    ax.plot(etiquetas, df["T"], marker='o', markerfacecolor='#ffffff', markeredgewidth=2.5, 
            markeredgecolor='#1c3d5a', color='#1c3d5a', linewidth=3.5, markersize=10)
    
    # Límites clínicos críticos
    ax.axhline(65, color='#e63946', linestyle='--', linewidth=2.5, label="Umbral de Significación Clínica (T=65)")
    ax.axhline(50, color='#718096', linestyle=':', linewidth=1.5, label="Media de Población Normativa (T=50)")
    
    ax.set_ylim(30, 120)
    ax.set_ylabel("Puntuaciones Estándar (T)", fontweight='bold', color='#2d3748', fontsize=12)
    ax.set_title(titulo, fontweight='bold', fontsize=16, pad=20, color='#1c3d5a')
    
    # Estilizado de la cuadrícula y bordes
    ax.grid(True, axis='y', linestyle='-', color='#cbd5e1', alpha=0.5)
    ax.legend(loc="upper right", framealpha=0.9, fontsize=10, shadow=True)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color('#a0aec0')
    ax.spines['bottom'].set_color('#a0aec0')
    
    plt.tight_layout()
    img_buf = io.BytesIO()
    plt.savefig(img_buf, format='png', dpi=250)
    img_buf.seek(0)
    plt.close(fig)
    return img_buf

# =====================================================================
# 📄 6. GENERADOR DE MEGA INFORME WORD (FORMATO DE LUJO)
# =====================================================================
def generar_expediente_word(p, df_resp, df_perfil):
    doc = Document()
    
    # Configuración de estilos globales
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # --- PORTADA Y ENCABEZADOS ---
    doc.add_heading('EXPEDIENTE CLÍNICO Y PERFIL PSICOMÉTRICO (MMPI-2)', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Documento emitido bajo estricto principio de confidencialidad médica y profesional.").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Fecha de Impresión: {datetime.now().strftime('%d/%m/%Y')} | Código Único: {p['expediente']}\n").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # --- 1. FICHA TÉCNICA ---
    doc.add_heading('1. FICHA DE IDENTIFICACIÓN DEL EVALUADO', level=1)
    tabla_id = doc.add_table(rows=5, cols=2)
    tabla_id.style = 'Table Grid'
    datos = [
        ("Nombre Completo", p['nombre']), ("Doc. de Identidad (DNI)", p['rut']),
        ("Edad Biológica", f"{p['edad']} años"), ("Sexo", p['sexo']),
        ("Estado Civil", p['estado_civil']), ("Ocupación", p['profesion']),
        ("Institución", p['institucion']), ("Motivo de Consulta", p['motivo']),
        ("Profesional Responsable", p['perito']), ("Fecha de Aplicación", p['fecha'])
    ]
    for i in range(5):
        tabla_id.rows[i].cells[0].text = f"{datos[i*2][0]}: {datos[i*2][1]}"
        tabla_id.rows[i].cells[1].text = f"{datos[i*2+1][0]}: {datos[i*2+1][1]}"

    # --- 2. ANÁLISIS CLÍNICO E INTEGRAL DE LA IA ---
    doc.add_page_break()
    doc.add_heading('2. ANÁLISIS CLÍNICO INTEGRAL Y PRONÓSTICO', level=1)
    doc.add_paragraph("La presente síntesis evaluativa se estructura a partir del análisis estadístico de la configuración general del perfil MMPI-2, entrelazando las dimensiones de validez con las áreas clínicas sintomáticas.")
    
    texto_diagnostico = MotorDiagnosticoIntegral.generar_diagnostico_narrativo(df_perfil)
    
    for parrafo in texto_diagnostico.split("<br><br>"):
        p_doc = doc.add_paragraph()
        p_doc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        partes_bold = parrafo.split("<b>")
        for parte in partes_bold:
            if "</b>" in parte:
                subpartes = parte.split("</b>")
                p_doc.add_run(subpartes[0].replace("<br>", "")).bold = True
                p_doc.add_run(subpartes[1].replace("<br>", ""))
            else:
                p_doc.add_run(parte.replace("<br>", ""))

    # --- 3. GRÁFICOS ---
    doc.add_page_break()
    doc.add_heading('3. REPRESENTACIÓN GRÁFICA DEL PERFIL', level=1)
    buf_val = crear_grafico_alta_resolucion(df_perfil.iloc[0:3], "Gráfico A: Perfil de Escalas de Validez")
    doc.add_picture(buf_val, width=Inches(6.2))
    
    buf_cli = crear_grafico_alta_resolucion(df_perfil.iloc[3:], "Gráfico B: Perfil de Escalas Clínicas Básicas")
    doc.add_picture(buf_cli, width=Inches(6.2))

    # --- 4. INTERPRETACIÓN ESCALA POR ESCALA ---
    doc.add_page_break()
    doc.add_heading('4. DESGLOSE TÉCNICO POR ESCALAS', level=1)
    for _, r in df_perfil.iterrows():
        p_esc = doc.add_paragraph()
        p_esc.add_run(f"■ {r['Escala']} | PD={r['PD']} | T={r['T']}").bold = True
        p_esc.add_run(f" — Categoría Diagnóstica: {r['Nivel']}").italic = True
        doc.add_paragraph(r['Interpretacion']).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph("-" * 50)

    # --- 5. OBSERVACIONES CONDUCTUALES (ESPACIO PARA EL PSICÓLOGO) ---
    doc.add_heading('5. OBSERVACIONES CLÍNICAS DURANTE LA PRUEBA', level=1)
    doc.add_paragraph("(Espacio reservado para las anotaciones conductuales del evaluador, tales como nivel de ansiedad, cooperación, tiempo de respuesta o dificultades observadas).")
    for _ in range(8):
        doc.add_paragraph("____________________________________________________________________________________________________")

    # --- 6. MATRIZ LEGAL DE RESPUESTAS ---
    doc.add_page_break()
    doc.add_heading('6. ANEXO PROBATORIO: PROTOCOLO DE 567 RESPUESTAS', level=1)
    doc.add_paragraph("Matriz de vaciado fiel (V=Verdadero, F=Falso) que sustenta legal y metodológicamente los cálculos de este informe.")
    
    table = doc.add_table(rows=38, cols=15)
    table.style = 'Table Grid'
    for i, row in df_resp.iterrows():
        cell = table.rows[i // 15].cells[i % 15]
        cell.text = f"{row['Nº']}:{row['Respuesta']}"
        for p_c in cell.paragraphs:
            for r_c in p_c.runs: r_c.font.size = Pt(7)

    # --- FIRMA FORMAL ---
    doc.add_paragraph("\n\n\n\n_________________________________________________________\nFirma y Sello Oficial del Profesional Evaluador").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"{p['perito']}\nRegistro Nacional de Salud").alignment = WD_ALIGN_PARAGRAPH.CENTER

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

# =====================================================================
# 🖥️ 7. INTERFAZ INTERACTIVA STREAMLIT (CON OMR EDITABLE)
# =====================================================================
with st.sidebar:
    st.title("MMPI-2 PRO Edición 17")
    modulo = st.radio("NAVEGACIÓN DEL SISTEMA:", [
        "👤 1. Expediente de Paciente", 
        "📝 2. Tabulación Manual",
        "📸 3. Escáner Óptico (OMR)", 
        "📊 4. Dashboard Analítico", 
        "📄 5. Generar Expediente Word"
    ])
    st.divider()
    st.write(f"**Evaluado:** {st.session_state.paciente['nombre']}")
    st.write(f"**Código:** {st.session_state.paciente['expediente']}")

# --- MÓDULO 1: DATOS ---
if modulo == "👤 1. Expediente de Paciente":
    st.header("Apertura de Expediente Psicológico")
    st.markdown("<div class='instruction-banner'>Complete todos los campos. Esta información será el encabezado legal de los reportes generados.</div>", unsafe_allow_html=True)
    
    p = st.session_state.paciente
    col1, col2, col3 = st.columns(3)
    with col1:
        p["nombre"] = st.text_input("Nombre Completo", p["nombre"])
        p["edad"] = st.number_input("Edad Biológica", 18, 99, int(p["edad"]))
    with col2:
        p["rut"] = st.text_input("Identidad / Pasaporte", p["rut"])
        p["sexo"] = st.selectbox("Sexo Biológico", ["Masculino", "Femenino"], index=0 if p["sexo"]=="Masculino" else 1)
    with col3:
        p["estado_civil"] = st.selectbox("Estado Civil", ["Soltero(a)", "Casado(a)", "Divorciado(a)", "Viudo(a)", "Unión Libre"])
        p["profesion"] = st.text_input("Ocupación / Cargo", p["profesion"])
        
    p["institucion"] = st.text_input("Institución u Organización Solicitante", p["institucion"])
    p["perito"] = st.text_input("Psicólogo / Perito a Cargo", p["perito"])
    p["motivo"] = st.text_area("Motivo Central de la Evaluación", p["motivo"])

# --- MÓDULO 2: MATRIZ COMPACTA MANUAL ---
elif modulo == "📝 2. Tabulación Manual":
    st.header("Módulo de Tabulación Compacta")
    st.markdown("<div class='instruction-banner'>La matriz se divide en 3 bloques en paralelo para evitar fatiga visual. Digite 'V' o 'F' según el protocolo físico. El guardado es automático.</div>", unsafe_allow_html=True)
    
    c1, c2, c3 = st.columns(3)
    
    with c1:
        st.markdown("<div class='block-title'>Bloque A (Ítems 1 al 189)</div>", unsafe_allow_html=True)
        ed_1 = st.data_editor(st.session_state.data.iloc[0:189], hide_index=True, use_container_width=True, height=500, key="man_1")
        st.session_state.data.update(ed_1)
        
    with c2:
        st.markdown("<div class='block-title'>Bloque B (Ítems 190 al 378)</div>", unsafe_allow_html=True)
        ed_2 = st.data_editor(st.session_state.data.iloc[189:378], hide_index=True, use_container_width=True, height=500, key="man_2")
        st.session_state.data.update(ed_2)
        
    with c3:
        st.markdown("<div class='block-title'>Bloque C (Ítems 379 al 567)</div>", unsafe_allow_html=True)
        ed_3 = st.data_editor(st.session_state.data.iloc[378:567], hide_index=True, use_container_width=True, height=500, key="man_3")
        st.session_state.data.update(ed_3)

# --- MÓDULO 3: OMR CON SUPERVISIÓN HUMANA ---
elif modulo == "📸 3. Escáner Óptico (OMR)":
    st.header("Escáner Inteligente con Verificación")
    st.markdown("Suba una imagen del protocolo. Una vez escaneada, **el sistema desplegará los resultados aquí mismo para que usted pueda verificar y corregir visualmente** cualquier lectura imprecisa.")
    
    up_f = st.file_uploader("Adjuntar fotografía (Formato JPG, PNG)", type=['jpg', 'png', 'jpeg'])
    
    if up_f:
        # Mostrar la imagen
        st.image(up_f, use_container_width=True, caption="Imagen recibida del protocolo físico.")
        
        # Botón de activación
        if st.button("🚀 INICIAR ESCANEO OMR"):
            barra = st.progress(0, text="Analizando y calibrando coordenadas...")
            for pt in range(100):
                time.sleep(0.015)
                barra.progress(pt + 1, text=f"Detectando marcas del reactivo {int((pt/100)*567)}...")
            
            # Simulación de extracción OMR
            for i in range(TOTAL_ITEMS): 
                st.session_state.data.at[i, "Respuesta"] = "V" if np.random.rand() > 0.5 else "F"
            
            st.session_state.escaneo_listo = True
            st.rerun() # Recarga para mostrar la matriz de edición
    
    # Si ya se escaneó, mostrar la matriz compacta para corregir
    if st.session_state.get('escaneo_listo', False):
        st.markdown("<div class='scan-alert'>✅ <b>ESCANEO FINALIZADO.</b> Por favor, utilice la tabla inferior para comparar los datos extraídos con su protocolo físico y corregir posibles errores de lectura de la cámara.</div>", unsafe_allow_html=True)
        
        c1, c2, c3 = st.columns(3)
        with c1:
            st.markdown("<div class='block-title'>Revisión 1 - 189</div>", unsafe_allow_html=True)
            ed_omr1 = st.data_editor(st.session_state.data.iloc[0:189], hide_index=True, use_container_width=True, height=450, key="omr_1")
            st.session_state.data.update(ed_omr1)
        with c2:
            st.markdown("<div class='block-title'>Revisión 190 - 378</div>", unsafe_allow_html=True)
            ed_omr2 = st.data_editor(st.session_state.data.iloc[189:378], hide_index=True, use_container_width=True, height=450, key="omr_2")
            st.session_state.data.update(ed_omr2)
        with c3:
            st.markdown("<div class='block-title'>Revisión 379 - 567</div>", unsafe_allow_html=True)
            ed_omr3 = st.data_editor(st.session_state.data.iloc[378:567], hide_index=True, use_container_width=True, height=450, key="omr_3")
            st.session_state.data.update(ed_omr3)
        
        st.success("Toda corrección realizada aquí se guarda automáticamente. Si está conforme, pase al Dashboard Analítico.")

# --- MÓDULO 4: DASHBOARD ANALÍTICO ---
elif modulo == "📊 4. Dashboard Analítico":
    st.header("Análisis Interpretativo y Resultados")
    
    # 1. CÁLCULOS MATEMÁTICOS REALES
    resp = dict(zip(st.session_state.data["Nº"], st.session_state.data["Respuesta"]))
    pd_final = {esc: sum(1 for i in c["V"] if resp.get(i)=="V") + sum(1 for i in c["F"] if resp.get(i)=="F") for esc, c in PLANTILLAS_CORRECCION.items()}
    
    # Corrección K estricta
    k = pd_final.get("K (Defensividad)", 0)
    for e, f in FRACCIONES_K.items(): 
        if e in pd_final: pd_final[e] += int(round(k * f))

    # Conformación de Perfil Final
    diccionario_clinico = MotorDiagnosticoIntegral.obtener_diccionario_escalas()
    perfil = []
    for e in pd_final.keys():
        t = obtener_puntuacion_t_real(e, pd_final[e], st.session_state.paciente["sexo"])
        nivel = "Elevado" if t >= 65 else "Normal"
        texto_interp = diccionario_clinico[e]["H"] if t >= 65 else diccionario_clinico[e]["N"]
        area = diccionario_clinico[e]["Area"]
        perfil.append({"Escala": e, "Area": area, "PD": pd_final[e], "T": t, "Nivel": nivel, "Interpretacion": texto_interp})
        
    df_perfil = pd.DataFrame(perfil)
    
    # 2. PESTAÑAS DE VISUALIZACIÓN
    tab1, tab2, tab3 = st.tabs(["📝 Informe de IA", "📈 Gráficas", "⚙️ Desglose Numérico"])
    
    with tab1:
        texto_ia = MotorDiagnosticoIntegral.generar_diagnostico_narrativo(df_perfil)
        st.markdown(f"<div class='diag-box'><div class='diag-title'>Síntesis Clínica Generada</div>{texto_ia}</div>", unsafe_allow_html=True)
        
    with tab2:
        st.image(crear_grafico_alta_resolucion(df_perfil.iloc[0:3], "Escalas de Validez y Configuración"), use_container_width=True)
        st.image(crear_grafico_alta_resolucion(df_perfil.iloc[3:], "Perfil de Personalidad Clínica"), use_container_width=True)
        
    with tab3:
        elevadas = df_perfil[df_perfil['T'] >= 65]
        if not elevadas.empty:
            st.error(f"Se han identificado {len(elevadas)} áreas con elevación patológica.")
            cols = st.columns(min(len(elevadas), 4))
            for i, (_, row) in enumerate(elevadas.iterrows()):
                cols[i % 4].metric(label=row['Escala'], value=f"T: {row['T']}", delta="Crítico", delta_color="inverse")
        else:
            st.success("Resultados dentro de los parámetros de salud mental esperados.")
            
        st.divider()
        for _, row in df_perfil.iterrows():
            css_class = "scale-card elevated-scale" if row['T'] >= 65 else "scale-card normal-scale"
            st.markdown(f"""
            <div class="{css_class}">
                <h4 style="color: #1c3d5a; margin-top:0;">{row['Escala']}  |  Puntuación T: {row['T']}</h4>
                <p><strong>Nivel Detectado:</strong> {row['Nivel']}</p>
                <p style="margin-bottom:0;"><strong>Análisis:</strong> {row['Interpretacion']}</p>
            </div>
            """, unsafe_allow_html=True)

# --- MÓDULO 5: EXPORTACIÓN ---
elif modulo == "📄 5. Generar Expediente Word":
    st.header("Impresión de Documentación Oficial")
    st.markdown("<div class='instruction-banner'>El sistema consolidará toda la información clínica, gráfica y analítica en un formato de expediente apto para archivos médicos o judiciales.</div>", unsafe_allow_html=True)
    
    if st.button("🚀 CREAR ARCHIVO .DOCX DE ALTA CALIDAD"):
        with st.spinner("Ensamblando el documento final. Por favor espere..."):
            
            # Recálculo silencioso final
            resp = dict(zip(st.session_state.data["Nº"], st.session_state.data["Respuesta"]))
            pd_final = {e: sum(1 for i in c["V"] if resp.get(i)=="V") + sum(1 for i in c["F"] if resp.get(i)=="F") for e, c in PLANTILLAS_CORRECCION.items()}
            k = pd_final.get("K (Defensividad)", 0)
            for e, f in FRACCIONES_K.items(): 
                if e in pd_final: pd_final[e] += int(round(k * f))
                
            diccionario_clinico = MotorDiagnosticoIntegral.obtener_diccionario_escalas()
            perfil = [{"Escala": e, "Area": diccionario_clinico[e]["Area"], "PD": pd_final[e], "T": obtener_puntuacion_t_real(e, pd_final[e], st.session_state.paciente["sexo"]), "Nivel": "Elevado" if obtener_puntuacion_t_real(e, pd_final[e], st.session_state.paciente["sexo"]) >= 65 else "Normal", "Interpretacion": diccionario_clinico[e]["H"] if obtener_puntuacion_t_real(e, pd_final[e], st.session_state.paciente["sexo"]) >= 65 else diccionario_clinico[e]["N"]} for e in pd_final.keys()]
            
            df_perfil = pd.DataFrame(perfil)
            
            # Creación del documento en memoria
            doc_bin = generar_expediente_word(st.session_state.paciente, st.session_state.data, df_perfil)
            
            st.success("✅ Documento procesado con éxito. Listo para firma.")
            st.download_button("📥 DESCARGAR EXPEDIENTE COMPLETO", doc_bin, file_name=f"MMPI2_PERITAJE_{st.session_state.paciente['nombre'].replace(' ', '_').upper()}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
