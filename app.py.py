import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import matplotlib.pyplot as plt
import math

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="Sistema Clínico MMPI-2", layout="wide")

# --- CONTROL DE SESIÓN ---
if 'revision' not in st.session_state:
    st.session_state.revision = 0
    st.session_state.respuestas = {i: None for i in range(1, 568)}
    st.session_state.pagina_actual = 0

def reset_form():
    st.session_state.revision += 1
    st.session_state.respuestas = {i: None for i in range(1, 568)}
    st.session_state.pagina_actual = 0
    st.rerun()

# =====================================================================
# 🛑 ZONA 1: ÍTEMS DEL CUESTIONARIO (Pega aquí tu manual) 🛑
# =====================================================================
PREGUNTAS = [
    "Me gustan las revistas de mecánica.", 
    "Tengo buen apetito.", 
    "Me despierto fresco y descansado casi todas las mañanas.",
    "Creo que me gustaría el trabajo de bibliotecario.",
    "El ruido me despierta fácilmente."
]
# Relleno temporal para que el código funcione hasta que pegues las 567
while len(PREGUNTAS) < 567:
    PREGUNTAS.append(f"Ítem reservado número {len(PREGUNTAS) + 1} del manual oficial.")

# =====================================================================
# 🛑 ZONA 2: CLAVES DE CORRECCIÓN (Copia de tu Word) 🛑
# =====================================================================
CLAVES = {
    "L (Mentira)": {"V": [], "F": [16, 29, 41, 51, 77, 93, 102, 107, 123, 139, 153, 183, 203, 232, 260]},
    "F (Incoherencia)": {"V": [18, 24, 30, 36, 42, 48, 54, 60, 66, 72, 84, 96, 114, 138, 144, 150, 156, 162, 168, 180, 198, 216, 228, 234, 240, 246, 252, 258, 264, 270, 282, 288, 294, 300, 306, 312, 324, 336, 349, 355, 361], "F": [6, 12, 78, 90, 102, 108, 120, 126, 132, 174, 186, 192, 204, 210, 222, 276, 318, 330, 343]},
    "K (Defensividad)": {"V": [83], "F": [29, 37, 58, 76, 110, 116, 122, 127, 130, 136, 148, 157, 158, 167, 171, 196, 213, 238, 240, 257, 258, 267, 281, 290, 300, 316, 319, 332, 338, 346]},
    "1 Hs (Hipocondriasis)": {"V": [11, 18, 28, 39, 59], "F": [2, 3, 9, 10, 20]}, # EJEMPLO: RELLENA CON EL MANUAL
    "2 D (Depresión)": {"V": [], "F": []},       
    "3 Hy (Histeria)": {"V": [], "F": []},
    "4 Pd (Desviación Psicopática)": {"V": [], "F": []},
    "5 Mf (Masculinidad-Feminidad)": {"V": [], "F": []}, 
    "6 Pa (Paranoia)": {"V": [], "F": []},
    "7 Pt (Psicastenia)": {"V": [11, 16, 23, 31, 38, 56, 65, 73, 82, 89, 94, 130, 147, 170, 175, 196, 218, 242, 273, 275, 277, 285, 289, 301, 302, 304, 308, 309, 310, 313, 316, 317, 320, 325, 326, 327, 328, 329, 331], "F": [3, 9, 33, 109, 140, 165, 174, 293, 321]},
    "8 Sc (Esquizofrenia)": {"V": [], "F": []},
    "9 Ma (Hipomanía)": {"V": [], "F": []},
    "0 Si (Introversión Social)": {"V": [], "F": []}
}

# =====================================================================
# 🛑 ZONA 3: BAREMOS T Y DESCRIPCIONES CLINICAS (De tus Excels) 🛑
# =====================================================================
def obtener_t(sexo, escala, puntaje):
    # EJEMPLO DE TABLA BAREMO (DEBES LLENARLA CON TUS EXCELS)
    baremos_ejemplo = {
        "Varón": {
            "1 Hs (Hipocondriasis)": {0: 30, 10: 50, 20: 70, 30: 90},
            # Añadir todas las puntuaciones reales aquí...
        },
        "Mujer": {
            "1 Hs (Hipocondriasis)": {0: 32, 10: 52, 20: 75, 30: 95},
            # Añadir todas las puntuaciones reales aquí...
        }
    }
    # Lógica para retornar T. Si no encuentra la exacta, simula una (QUITAR SIMULACIÓN LUEGO)
    try:
        return baremos_ejemplo[sexo][escala][puntaje]
    except KeyError:
        return min(puntaje * 2 + 30, 120) # SIMULACIÓN: ELIMINAR CUANDO LLENES LA TABLA REAL

def obtener_interpretacion(escala, t_score):
    if t_score >= 65:
        descripciones = {
            "L (Mentira)": "Cuadro defensivo. Intenta mostrar una imagen de perfección moral. Posible negación de problemas.",
            "F (Incoherencia)": "Exageración de síntomas, fingimiento (fake bad), o patología severa (procesos psicóticos).",
            "K (Defensividad)": "Fingir buena imagen. Marcada defensividad clínica. Resistencia a la evaluación.",
            "1 Hs (Hipocondriasis)": "Preocupaciones somáticas graves. Constreñido, inmovilizado por quejas físicas. Exigente.",
            "2 D (Depresión)": "Depresión moderada/severa. Insatisfacción, pesimismo, falta de energía y problemas de sueño.",
            "3 Hy (Histeria)": "Síntomas somáticos ante el estrés. Inmadurez, demanda de atención, sugestionabilidad.",
            "4 Pd (Desviación Psicopática)": "Rebeldía, impulsividad, problemas con la autoridad, relaciones superficiales, baja tolerancia a la frustración.",
            "5 Mf (Masculinidad-Feminidad)": "Patrón atípico de intereses respecto al rol de género tradicional.",
            "6 Pa (Paranoia)": "Predisposición paranoide. Excesivamente sensible, suspicaz, resentido y desconfiado.",
            "7 Pt (Psicastenia)": "Ansiedad elevada, obsesiones, compulsiones, culpa, tensión constante y miedos.",
            "8 Sc (Esquizofrenia)": "Aislamiento social, confusión mental, creencias inusuales, posible desconexión de la realidad.",
            "9 Ma (Hipomanía)": "Aceleración psicomotriz, irritabilidad, grandiosidad, impulsividad, baja necesidad de sueño.",
            "0 Si (Introversión Social)": "Aislamiento social marcado, evitación de contacto, timidez extrema."
        }
        return descripciones.get(escala, "Elevación clínica significativa en esta escala.")
    else:
        return "Dentro de los límites normativos (No significativo)."

# =====================================================================
# LÓGICA DE INTERFAZ Y PAGINACIÓN
# =====================================================================

ITEMS_POR_PAGINA = 50
total_paginas = math.ceil(567 / ITEMS_POR_PAGINA)

with st.sidebar:
    st.header("📋 Datos Clínicos")
    nombre = st.text_input("Nombre Completo")
    edad = st.number_input("Edad", 18, 99, 25)
    sexo = st.selectbox("Sexo Biológico", ["Varón", "Mujer"])
    
    st.divider()
    progreso = sum(1 for r in st.session_state.respuestas.values() if r is not None) / 567
    st.progress(progreso)
    st.write(f"Avance: {int(progreso * 100)}% (Pág. {st.session_state.pagina_actual + 1} de {total_paginas})")
    
    if st.button("🗑️ Reiniciar Evaluación"): reset_form()

st.title("MMPI-2: Inventario Multifásico de Personalidad")
st.info("Responda Verdadero o Falso a cada afirmación. Solo seleccione 'No sé' si es absolutamente indispensable.")

# CUESTIONARIO PAGINADO
inicio = st.session_state.pagina_actual * ITEMS_POR_PAGINA
fin = min(inicio + ITEMS_POR_PAGINA, 567)

for i in range(inicio + 1, fin + 1):
    st.session_state.respuestas[i] = st.radio(
        f"**{i}. {PREGUNTAS[i-1]}**", 
        ["Verdadero", "Falso", "No sé (Blanco)"], 
        horizontal=True, 
        index=["Verdadero", "Falso", "No sé (Blanco)"].index(st.session_state.respuestas[i]) if st.session_state.respuestas[i] else None,
        key=f"q_{i}"
    )

st.divider()

# NAVEGACIÓN Y CÁLCULO
col1, col2, col3 = st.columns([1, 2, 1])
with col1:
    if st.session_state.pagina_actual > 0:
        if st.button("⬅️ Anterior"):
            st.session_state.pagina_actual -= 1
            st.rerun()
with col3:
    if st.session_state.pagina_actual < total_paginas - 1:
        if st.button("Siguiente ➡️"):
            st.session_state.pagina_actual += 1
            st.rerun()
    else:
        if st.button("📈 CALCULAR PERFIL Y GENERAR INFORME"):
            # Validación: Verificar ítems en blanco (?)
            blancos = sum(1 for v in st.session_state.respuestas.values() if v is None or v == "No sé (Blanco)")
            
            if blancos > 30:
                st.error(f"⚠️ El test es INVÁLIDO. Hay {blancos} ítems sin responder (Límite: 30).")
            else:
                # 1. CÁLCULO DE PUNTUACIÓN DIRECTA (PD)
                resultados = {}
                for escala, claves in CLAVES.items():
                    pd = 0
                    for item in claves["V"]:
                        if st.session_state.respuestas[item] == "Verdadero": pd += 1
                    for item in claves["F"]:
                        if st.session_state.respuestas[item] == "Falso": pd += 1
                    resultados[escala] = {"PD": pd, "PD_K": pd, "T": 0, "Interp": ""}

                # 2. MOTOR DE CORRECCIÓN K (Matemática del Manual MMPI-2)
                valor_k = resultados["K (Defensividad)"]["PD"]
                
                if "1 Hs (Hipocondriasis)" in resultados:
                    resultados["1 Hs (Hipocondriasis)"]["PD_K"] = round(resultados["1 Hs (Hipocondriasis)"]["PD"] + (0.5 * valor_k))
                if "4 Pd (Desviación Psicopática)" in resultados:
                    resultados["4 Pd (Desviación Psicopática)"]["PD_K"] = round(resultados["4 Pd (Desviación Psicopática)"]["PD"] + (0.4 * valor_k))
                if "7 Pt (Psicastenia)" in resultados:
                    resultados["7 Pt (Psicastenia)"]["PD_K"] = round(resultados["7 Pt (Psicastenia)"]["PD"] + (1.0 * valor_k))
                if "8 Sc (Esquizofrenia)" in resultados:
                    resultados["8 Sc (Esquizofrenia)"]["PD_K"] = round(resultados["8 Sc (Esquizofrenia)"]["PD"] + (1.0 * valor_k))
                if "9 Ma (Hipomanía)" in resultados:
                    resultados["9 Ma (Hipomanía)"]["PD_K"] = round(resultados["9 Ma (Hipomanía)"]["PD"] + (0.2 * valor_k))

                # 3. CONVERSIÓN A PUNTUACIÓN T Y EXTRACCIÓN DE DIAGNÓSTICO
                for escala in resultados.keys():
                    puntaje_base = resultados[escala]["PD_K"]
                    t_score = obtener_t(sexo, escala, puntaje_base)
                    resultados[escala]["T"] = t_score
                    resultados[escala]["Interp"] = obtener_interpretacion(escala, t_score)

                # --- VISUALIZACIÓN EN PANTALLA ---
                st.success("Perfil Generado Correctamente")
                
                df = pd.DataFrame([
                    {"Escala": k, "PD": v["PD"], "PD+K": v["PD_K"], "Puntuación T": v["T"], "Interpretación": v["Interp"]} 
                    for k, v in resultados.items()
                ])

                # Gráfico de Líneas (Perfil MMPI Clásico)
                fig, ax = plt.subplots(figsize=(12, 6))
                
                # Separar validez y clínicas para el gráfico si se desea, aquí graficamos todas juntas
                ax.plot(df['Escala'], df['Puntuación T'], marker='o', linewidth=2, color='darkblue')
                ax.axhline(y=65, color='red', linestyle='--', linewidth=2, label='Significancia Clínica (T=65)')
                ax.axhline(y=50, color='gray', linestyle='-', alpha=0.5, label='Media Normativa (T=50)')
                
                # Sombreado de área clínica
                ax.fill_between(df['Escala'], 65, 120, color='red', alpha=0.1)
                
                ax.set_ylim(20, 120)
                ax.set_ylabel('Puntuación T')
                ax.set_title(f'Perfil Clínico MMPI-2: {nombre}')
                plt.xticks(rotation=45, ha='right')
                ax.legend()
                ax.grid(True, linestyle=':', alpha=0.6)
                
                st.pyplot(fig)
                st.dataframe(df[['Escala', 'PD', 'PD+K', 'Puntuación T']])

                # --- GENERACIÓN DE REPORTE WORD ---
                doc = Document()
                sec = doc.sections[0]
                sec.top_margin = sec.bottom_margin = Inches(0.5)
                sec.left_margin = sec.right_margin = Inches(0.5)

                doc.add_heading('REPORTE PSICOLÓGICO: MMPI-2', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                doc.add_heading('1. Datos de Identificación', level=1)
                doc.add_paragraph(f"Nombre: {nombre}\nEdad: {edad} años\nSexo de Baremos: {sexo}\nÍtems omitidos (?): {blancos}")

                # Gráfico
                doc.add_heading('2. Perfil Gráfico de Escalas', level=1)
                img_b = BytesIO()
                plt.savefig(img_b, format='png', bbox_inches='tight')
                doc.add_picture(img_b, width=Inches(6.5))

                # Tabla de Puntuaciones
                doc.add_heading('3. Resumen de Puntuaciones', level=1)
                t = doc.add_table(rows=1, cols=4)
                t.style = 'Table Grid'
                hdr = t.rows[0].cells
                hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'Escala', 'PD', 'PD+K', 'T'
                for index, row in df.iterrows():
                    row_cells = t.add_row().cells
                    row_cells[0].text = str(row['Escala'])
                    row_cells[1].text = str(row['PD'])
                    row_cells[2].text = str(row['PD+K'])
                    row_cells[3].text = str(row['Puntuación T'])

                # Interpretación Clínica Automática
                doc.add_page_break()
                doc.add_heading('4. Interpretación Clínica Descriptiva', level=1)
                
                escalas_altas = df[df['Puntuación T'] >= 65]
                
                if escalas_altas.empty:
                    doc.add_paragraph("No se registran elevaciones clínicamente significativas (T ≥ 65) en las escalas básicas. Perfil dentro de los límites normativos.")
                else:
                    doc.add_paragraph("Se detectaron las siguientes elevaciones clínicamente significativas (T ≥ 65):")
                    for index, row in escalas_altas.iterrows():
                        p = doc.add_paragraph()
                        p.add_run(f"■ {row['Escala']} (T = {row['Puntuación T']}): ").bold = True
                        p.add_run(row['Interpretación'])

                w_buf = BytesIO()
                doc.save(w_buf)
                st.download_button("📥 DESCARGAR REPORTE PROFESIONAL", w_buf.getvalue(), f"MMPI2_Reporte_{nombre}.docx")